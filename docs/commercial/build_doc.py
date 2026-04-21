"""Build Wamply Business Case Word document for sales + marketing team.

Sections:
1. Cover
2. Executive Summary
3. Product
4. Business Model (plans + top-up)
5. Cost Structure
6. Margin Analysis per plan
7. Top-up Economics
8. Azure Enterprise Edition
9. Unit Economics & Scenarios
10. Go-to-Market
11. Roadmap
12. Appendix

Brand: Wamply navy (#1B2A4A) + teal (#0D9488).
Source data: ai-credits-plan.md v2 (2026-04-21).
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).parent
COVER_PNG = OUT / "wamply-cover.png"
LOGO_PNG = OUT / "wamply-logo.png"

# Brand palette
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
NAVY_DEEP = RGBColor(0x0B, 0x16, 0x28)
TEAL = RGBColor(0x0D, 0x94, 0x88)
TEAL_SOFT = RGBColor(0xCC, 0xFB, 0xF1)
GREEN = RGBColor(0x25, 0xD3, 0x66)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
GRAY_LIGHT = RGBColor(0x94, 0xA3, 0xB8)
GRAY_MUTED = RGBColor(0x64, 0x74, 0x8B)
INK = RGBColor(0x1B, 0x2A, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Shading helpers


def cell_shade(cell, hex_color: str) -> None:
    """Apply background color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, right=120, bottom=80, left=120) -> None:
    """Set cell internal padding in twips."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for m, v in [("top", top), ("right", right), ("bottom", bottom), ("left", left)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def style_header_cell(cell, text: str) -> None:
    cell_shade(cell, "1B2A4A")
    set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = WHITE


def style_data_cell(cell, text: str, bold: bool = False,
                    color: RGBColor = INK, shade: str | None = None,
                    align: str = "left") -> None:
    if shade:
        cell_shade(cell, shade)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = bold
    run.font.color.rgb = color


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.bold = True
    run.font.color.rgb = NAVY if level == 1 else INK
    run.font.size = Pt(20 if level == 1 else 14 if level == 2 else 12)
    # Outline level for TOC
    p_pr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level - 1))
    p_pr.append(outline)


def add_para(doc, text: str, size: int = 10, color: RGBColor = INK,
             italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.italic = italic
    run.font.bold = bold


def add_bullet(doc, text: str, size: int = 10) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = INK


def add_teal_divider(doc) -> None:
    """Horizontal teal line as a paragraph border."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0D9488")
    pbdr.append(bottom)
    p_pr.append(pbdr)


def add_callout(doc, text: str, shade_hex: str = "F0FDFA",
                border_hex: str = "0D9488") -> None:
    """1-cell teal-accent table used as callout box."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.columns[0].width = Cm(16)
    cell = t.cell(0, 0)
    cell.width = Cm(16)
    cell_shade(cell, shade_hex)
    set_cell_margins(cell, top=180, right=240, bottom=180, left=240)
    # Left accent border
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D1E7E5")
        tc_borders.append(b)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), border_hex)
    tc_borders.append(left)
    tc_pr.append(tc_borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = INK


def build_table(doc, headers: list[str], rows: list[list[str]],
                highlight_col: int | None = None,
                widths_cm: list[float] | None = None,
                highlight_rows: dict | None = None) -> None:
    """headers row styled navy, data rows alternating white/slate.
    highlight_col highlights a specific column in teal-soft.
    highlight_rows is dict {row_idx: hex_color} to highlight specific rows.
    """
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    if widths_cm:
        for i, w in enumerate(widths_cm):
            t.columns[i].width = Cm(w)
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    for i, h in enumerate(headers):
        style_header_cell(t.rows[0].cells[i], h)
    highlight_rows = highlight_rows or {}
    for r_idx, row in enumerate(rows):
        row_shade = highlight_rows.get(r_idx)
        for c_idx, txt in enumerate(row):
            if row_shade:
                shade = row_shade
            elif highlight_col is not None and c_idx == highlight_col:
                shade = "F0FDFA"
            else:
                shade = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
            bold = c_idx == 0 or (highlight_col is not None and c_idx == highlight_col)
            align = "right" if c_idx > 0 and any(ch.isdigit() for ch in txt) else "left"
            style_data_cell(t.rows[r_idx + 1].cells[c_idx], txt,
                            bold=bold, shade=shade, align=align)


def add_footer(doc) -> None:
    """Footer with page number + confidential note."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Wamply Business Case  ·  Confidenziale  ·  Uso interno team Sales & Marketing  ·  ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY_MUTED
    # Page number field
    run2 = p.add_run("Pag. ")
    run2.font.name = "Arial"
    run2.font.size = Pt(8)
    run2.font.color.rgb = GRAY_MUTED
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run2._r.append(fld_char1)
    run2._r.append(instr)
    run2._r.append(fld_char2)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

# A4 with 1" margins approximately
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

# Default style
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(10)
style.font.color.rgb = INK

add_footer(doc)

# ===== COVER PAGE =====
doc.add_picture(str(COVER_PNG), width=Cm(16))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Business Case")
run.font.name = "Arial"
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Margini, Costi e Strategia Go-to-Market")
run.font.name = "Arial"
run.font.size = Pt(18)
run.font.color.rgb = TEAL

add_teal_divider(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Documento interno per team Commerciale & Marketing")
run.font.name = "Arial"
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = GRAY_MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("Versione 1.0  ·  21 aprile 2026")
run.font.name = "Arial"
run.font.size = Pt(11)
run.font.color.rgb = GRAY_LIGHT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run("CONFIDENZIALE")
run.font.name = "Arial"
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = AMBER

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Preparato da Agent Engineering Studio")
run.font.name = "Arial"
run.font.size = Pt(10)
run.font.color.rgb = GRAY_MUTED

doc.add_page_break()

# ===== 1. EXECUTIVE SUMMARY =====
add_heading(doc, "1. Executive Summary", 1)

add_para(doc, "Wamply e' una piattaforma SaaS per campagne WhatsApp personalizzate tramite AI. Il prodotto automatizza il lavoro di copywriting, segmentazione e ottimizzazione di campagne di messaggistica massiva tramite l'integrazione di Claude (Anthropic) come agent intelligente.", 10)

add_para(doc, "A differenza dei broadcaster tradizionali (Mailchimp, Brevo) che inviano lo stesso template a tutti, Wamply genera un messaggio unico per ogni destinatario, in tempo reale, con tono e lingua appropriati al contesto.", 10)

add_heading(doc, "1.1 Proposta di valore", 2)
add_bullet(doc, "AI-first: l'AI e' parte integrante del prodotto, non un add-on")
add_bullet(doc, "Routing silente: l'utente paga solo in crediti, Wamply sceglie il modello AI ottimale per ogni operazione")
add_bullet(doc, "BYOK supportato: utenti avanzati portano la loro API key Claude, zero costi AI aggiuntivi")
add_bullet(doc, "Top-up on demand: pacchetti crediti one-shot per scalare senza cambiare piano")
add_bullet(doc, "Pricing trasparente: €49 / €149 / €399 per piani standard + Azure Enterprise Edition custom")

add_heading(doc, "1.2 Obiettivi commerciali", 2)
add_callout(
    doc,
    "Target Q1-Q2 2026: 50 clienti Professional + 10 Enterprise + 2 Azure Enterprise Edition. "
    "ARR proiettato: ~€210.000. Margine medio ponderato: 82%. "
    "Break-even piattaforma raggiunto con 70 clienti Pro o equivalenti."
)

add_heading(doc, "1.3 Riepilogo margini", 2)
build_table(
    doc,
    headers=["Piano", "Prezzo", "Costo AI", "Altri costi", "Margine lordo", "Margine %"],
    rows=[
        ["Starter", "€49", "€0 (BYOK)", "€5", "€44", "90%"],
        ["Professional", "€149", "€4.80", "€6.50", "€138", "92%"],
        ["Enterprise", "€399", "€48", "€11", "€340", "85%"],
        ["Azure Enterprise", "€1.500+", "€120+", "€280+", "€1.100+", "65-75%"],
    ],
    widths_cm=[3.5, 2.0, 2.5, 2.5, 3.0, 2.5],
    highlight_rows={2: "F0FDFA"},
)
add_para(doc, "I margini sono lordi sul piano mensile (escluso costo Twilio WhatsApp, gestito come pass-through). Vedi sezione 6 per calcoli dettagliati.", 9, italic=True, color=GRAY_MUTED)

doc.add_page_break()

# ===== 2. IL PRODOTTO =====
add_heading(doc, "2. Il Prodotto", 1)
add_para(doc, "Wamply integra tre componenti chiave che lavorano insieme per automatizzare campagne WhatsApp di alto livello.")

add_heading(doc, "2.1 Componenti tecnici", 2)
add_bullet(doc, "Frontend: Next.js 15 + React 19 - dashboard utente e admin")
add_bullet(doc, "Backend: Python FastAPI - gestione piani, fatturazione Stripe, API templates")
add_bullet(doc, "Agent AI: Microsoft Agent Framework + Claude (Sonnet / Opus / Haiku con routing automatico)")
add_bullet(doc, "Infrastruttura: Supabase (PostgreSQL + Auth) + Redis + Twilio WhatsApp API")

add_heading(doc, "2.2 Differenziali competitivi", 2)
add_bullet(doc, "Prima piattaforma italiana con AI routing automatico: semplicita' d'uso + ottimizzazione costi")
add_bullet(doc, "Personalizzazione vera - non template con variabili, ma messaggio unico per destinatario")
add_bullet(doc, "Supporto nativo BYOK - per aziende con contratti Anthropic diretti")
add_bullet(doc, "Compliance-ready: PCI DSS Level 1, 3D Secure 2, TLS 1.3, dati EU")

doc.add_page_break()

# ===== 3. BUSINESS MODEL =====
add_heading(doc, "3. Business Model", 1)

add_heading(doc, "3.1 Piani principali", 2)
build_table(
    doc,
    headers=["Piano", "Prezzo/mese", "Target cliente", "Crediti AI inclusi", "BYOK"],
    rows=[
        ["Free Trial 14gg", "€0", "Lead qualificati", "200 (eredita Pro)", "No"],
        ["Starter", "€49", "Freelance, PMI piccole", "0 - solo BYOK", "Si"],
        ["Professional", "€149", "PMI mid-market", "200/mese", "Si"],
        ["Enterprise", "€399", "Aziende strutturate", "1.500/mese", "Illimitato"],
        ["Azure Enterprise", "da €1.500", "Corporate, PA, regolamentati", "Custom", "Illimitato"],
    ],
    widths_cm=[3.5, 2.2, 4.0, 3.5, 2.3],
    highlight_rows={3: "F0FDFA"},
)
add_para(doc, "Il piano consigliato e' Professional (badge \"Il piu' scelto\" sul sito). Enterprise e' il vettore di espansione. Azure Enterprise e' il lift premium per mercati regolamentati.", 9, italic=True, color=GRAY_MUTED)

add_heading(doc, "3.2 Pacchetti Top-up crediti", 2)
add_para(doc, "Disponibili per utenti con piano attivo. Crediti validi 12 mesi dall'ultimo acquisto. Consumo \"piano-first\": prima i crediti del piano mensile, poi i top-up. Incentivo volume: sconto progressivo per pacchetti piu' grandi.")

build_table(
    doc,
    headers=["Pacchetto", "Crediti", "Prezzo", "€/credito", "Badge", "Target"],
    rows=[
        ["Small", "100", "€15", "€0,15", "-", "Prova / picco occasionale"],
        ["Medium", "500", "€59", "€0,118", "Piu' venduto", "Boost mensile Pro"],
        ["Large", "2.000", "€199", "€0,10", "Miglior prezzo", "Enterprise overflow"],
        ["XL", "10.000", "€799", "€0,08", "-", "Campagne massive"],
    ],
    widths_cm=[2.2, 2.0, 2.0, 2.3, 2.8, 4.5],
    highlight_rows={1: "F0FDFA"},
)

add_callout(
    doc,
    "Regola operativa: BYOK e utenti Free non possono acquistare top-up. "
    "BYOK = l'utente paga Anthropic direttamente, Wamply non conteggia crediti. "
    "Free = deve prima scegliere un piano. Pattern standard in linea con OpenAI/altri SaaS."
)

doc.add_page_break()

# ===== 4. COST STRUCTURE =====
add_heading(doc, "4. Struttura dei Costi", 1)

add_heading(doc, "4.1 Costi variabili AI (Anthropic)", 2)
add_para(doc, "Il backend applica routing automatico al modello ottimale per ogni operazione. Il costo medio pesato e' ~$0,026 per credito (€0,024 al cambio corrente).")

build_table(
    doc,
    headers=["Operazione AI", "Crediti", "Modello (invisibile utente)", "Costo stimato"],
    rows=[
        ["Chat turn standard", "1.0", "Sonnet", "$0.020"],
        ["Chat turn con tool-use", "2.0", "Sonnet", "$0.040"],
        ["Chat turn con pianificazione", "3.0", "Opus", "$0.080"],
        ["Template generate", "2.0", "Sonnet", "$0.040"],
        ["Template improve (3 varianti)", "3.0", "Sonnet", "$0.060"],
        ["Template compliance check", "3.0", "Opus", "$0.090"],
        ["Template translate", "1.0", "Haiku", "$0.005"],
        ["Personalize per messaggio", "0.5", "Haiku", "$0.003"],
        ["Batch bulk personalize", "0.5", "Haiku", "$0.003"],
        ["Planner campagna (endpoint dedicato)", "5.0", "Opus", "$0.150"],
    ],
    widths_cm=[6.0, 2.0, 5.0, 3.0],
)
add_para(doc, "Il mix reale d'uso e' ~60% Sonnet, 25% Haiku, 15% Opus. Con 200 crediti pro-rata, costo Anthropic stimato: $5,20 (~€4,80). Con 1.500 crediti Enterprise: $52 (~€48).", 9, italic=True, color=GRAY_MUTED)

add_heading(doc, "4.2 Altri costi variabili per cliente", 2)
build_table(
    doc,
    headers=["Voce di costo", "Importo", "Note"],
    rows=[
        ["Twilio WhatsApp API", "$0,005-0,12 per msg", "Pass-through al cliente (markup 0-10%)"],
        ["Stripe fees EU", "1,5% + €0,25", "Per transazione (carte UE)"],
        ["Email SMTP Aruba", "€0,002 per email", "Reminder trial + transactional"],
        ["Monitoring / logging", "~€0,30/utente/mese", "Jaeger, logs, Redis InsightRate"],
    ],
    widths_cm=[5.0, 4.0, 7.0],
)

add_heading(doc, "4.3 Costi fissi di piattaforma", 2)
add_para(doc, "Costi che non scalano direttamente col numero di clienti, spalmati sulla base utenti.")

build_table(
    doc,
    headers=["Voce", "Costo mensile", "Scalabilita'"],
    rows=[
        ["Hosting Kubernetes (backend + agent + frontend)", "€200-500", "Fino a 500 utenti"],
        ["Supabase Postgres + GoTrue + Redis", "€100-300", "Fino a 1000 utenti"],
        ["Kong Gateway + CDN + domini + SSL", "€50", "Fisso"],
        ["Monitoring (Sentry, Jaeger cloud)", "€50-100", "Fisso"],
        ["Compliance + legal + contabilita'", "€300", "Fisso"],
        ["Customer support (part-time)", "€1.500-3.000", "Scala con Enterprise"],
        ["TOTALE piattaforma", "€2.200-4.250", ""],
    ],
    widths_cm=[8.0, 4.0, 4.0],
    highlight_rows={6: "FFF5E6"},
)

doc.add_page_break()

# ===== 5. MARGIN ANALYSIS =====
add_heading(doc, "5. Analisi Margini per Piano", 1)

# Starter
add_heading(doc, "5.1 Piano Starter (€49/mese)", 2)
add_para(doc, "Piano di ingresso BYOK-first. Il cliente usa la sua API key Claude, Wamply non conteggia crediti AI. Pensato per freelance, professionisti, PMI che hanno gia' familiarita' con Anthropic o vogliono controllo diretto dei costi AI.")

build_table(
    doc,
    headers=["Voce", "Valore", "Note"],
    rows=[
        ["Revenue mensile", "€49,00", "Pagamento Stripe ricorrente"],
        ["Costo AI (Anthropic)", "€0,00", "BYOK - paga il cliente"],
        ["Stripe fees (1,5% + €0,25)", "-€0,99", ""],
        ["Costo infrastruttura", "-€4,00", "Quota piatta spalmata"],
        ["Margine lordo", "€44,01", "89,8%"],
        ["Margine netto post-CAC", "€18-38", "CAC stimato €50-200 ammortizzato 12 mesi"],
    ],
    widths_cm=[6.5, 4.0, 5.5],
    highlight_rows={4: "F0FDFA", 5: "FFFBEB"},
)
add_callout(
    doc,
    "Insight vendita: Starter e' il piano con margine % piu' alto ma revenue assoluta bassa. "
    "Usato come \"gateway\" - il 40% degli Starter attivi passa a Professional entro 6 mesi "
    "quando serve l'AI system-managed per scalare."
)

# Professional
add_heading(doc, "5.2 Piano Professional (€149/mese)", 2)
add_para(doc, "Piano consigliato (badge \"Il piu' scelto\" sul sito). Include 200 crediti AI/mese con modello auto-selected. Target: marketing team di PMI mid-market che fanno 1-3 campagne/settimana.")

build_table(
    doc,
    headers=["Voce", "Valore", "Note"],
    rows=[
        ["Revenue mensile", "€149,00", "Pagamento Stripe ricorrente"],
        ["Costo AI (200 crediti @ $0.026)", "-€4,80", "Mix Sonnet 60% + Haiku 25% + Opus 15%"],
        ["Stripe fees", "-€2,49", ""],
        ["Costo infrastruttura", "-€4,00", ""],
        ["Margine lordo", "€137,71", "92,4%"],
        ["Margine netto post-CAC", "€85-120", "CAC €150-400 ammortizzato 18 mesi"],
    ],
    widths_cm=[6.5, 4.0, 5.5],
    highlight_rows={4: "F0FDFA", 5: "FFFBEB"},
)
add_callout(
    doc,
    "Insight vendita: il Professional ha il LTV/CAC piu' favorevole. Con retention 70% su 18 mesi, "
    "LTV ~€1.880, break-even CAC raggiunto entro 3 mesi. E' il pilastro della growth machine."
)

# Enterprise
add_heading(doc, "5.3 Piano Enterprise (€399/mese)", 2)
add_para(doc, "Per aziende strutturate con esigenze multi-team. 1.500 crediti/mese + BYOK illimitato + white-label + supporto dedicato. Target: aziende retail, servizi, finanza con 10+ addetti marketing.")

build_table(
    doc,
    headers=["Voce", "Valore", "Note"],
    rows=[
        ["Revenue mensile", "€399,00", "Pagamento Stripe ricorrente"],
        ["Costo AI (1.500 crediti @ $0.026)", "-€36,00", "Mix simile a Pro"],
        ["Stripe fees", "-€6,24", ""],
        ["Costo infrastruttura + support", "-€15,00", "Quota piu' alta: CSM dedicato part-time"],
        ["Margine lordo", "€341,76", "85,7%"],
        ["Margine netto post-CAC", "€200-280", "CAC outbound €1.500-3.000 ammortizzato 24 mesi"],
    ],
    widths_cm=[6.5, 4.0, 5.5],
    highlight_rows={4: "F0FDFA", 5: "FFFBEB"},
)
add_callout(
    doc,
    "Insight vendita: Enterprise ha revenue assoluta 2,7x Professional. LTV medio €8.140 (24 mesi retention 85%). "
    "Da negoziare sempre contratto annuale con sconto 10-15% per assicurare retention e pagamento anticipato."
)

# Free Trial
add_heading(doc, "5.4 Free Trial 14 giorni", 2)
add_para(doc, "Il trial eredita il piano Professional per 14 giorni. E' il principale driver di conversione del funnel di acquisizione. Nessuna carta richiesta al signup per ridurre attrito.")

build_table(
    doc,
    headers=["Voce", "Valore", "Note"],
    rows=[
        ["Revenue", "€0", "-"],
        ["Costo AI stimato (media 100 crediti/14gg)", "-€2,40", "I trial consumano meno della meta'"],
        ["Costo infrastruttura", "-€2,00", "14/30 giorni di fisso"],
        ["Costo per trial", "-€4,40", ""],
        ["Conversion rate atteso", "15-25%", "Basato su benchmark SaaS B2B"],
        ["Costo acquisizione per cliente pagante", "€18-29", "€4,40 diviso tasso conversione"],
    ],
    widths_cm=[6.5, 4.0, 5.5],
    highlight_rows={3: "FFF5E6", 5: "F0FDFA"},
)

doc.add_page_break()

# ===== 6. TOP-UP ECONOMICS =====
add_heading(doc, "6. Economia dei Top-up", 1)
add_para(doc, "I pacchetti top-up sono acquisti one-shot attivabili quando il cliente esaurisce i crediti mensili. Hanno margine inferiore rispetto al piano (non c'e' la \"rendita\" mensile) ma generano revenue aggiuntiva significativa.")

build_table(
    doc,
    headers=["Pacchetto", "Crediti", "Prezzo", "Costo AI (*)", "Stripe", "Margine €", "Margine %"],
    rows=[
        ["Small", "100", "€15", "-€2,40", "-€0,48", "€12,12", "80,8%"],
        ["Medium", "500", "€59", "-€12,00", "-€1,14", "€45,86", "77,7%"],
        ["Large", "2.000", "€199", "-€48,00", "-€3,24", "€147,76", "74,3%"],
        ["XL", "10.000", "€799", "-€240,00", "-€12,24", "€546,76", "68,4%"],
    ],
    widths_cm=[2.2, 2.0, 2.0, 2.5, 2.0, 2.5, 2.3],
    highlight_rows={1: "F0FDFA"},
)
add_para(doc, "(*) Costo AI calcolato come crediti × $0.026. Non include quota infra (fissa, non scala col top-up).", 9, italic=True, color=GRAY_MUTED)

add_heading(doc, "6.1 Proiezione revenue top-up", 2)
add_callout(
    doc,
    "Assumendo che il 20% dei clienti Professional acquisti 2 pacchetti Medium/anno: "
    "100 Pro x 20% x 2 x €59 = €2.360/anno di revenue top-up aggiuntiva. "
    "Con margine 78%: €1.840/anno di margine incrementale senza nuovi clienti."
)

add_heading(doc, "6.2 Leve commerciali top-up", 2)
add_bullet(doc, "Email warning automatica a 80% consumo crediti: CTA \"Ricarica crediti\" con pacchetto Medium evidenziato")
add_bullet(doc, "Banner in dashboard per utenti che hanno raggiunto il limite: CTA \"Ricarica\" + \"Passa a Enterprise\"")
add_bullet(doc, "Pacchetto Large posizionato come \"Miglior prezzo\" per stimolare l'upsell volume")
add_bullet(doc, "Pacchetto XL per eventi speciali (Black Friday, lancio prodotto) - remarketing mirato")

doc.add_page_break()

# ===== 7. AZURE ENTERPRISE EDITION =====
add_heading(doc, "7. Azure Enterprise Edition (premium)", 1)
add_para(doc, "L'Azure Enterprise Edition e' un'offerta dedicata per clienti che richiedono deployment su infrastruttura Microsoft Azure, SLA garantiti e conformita' a standard regolamentari stringenti (banche, sanita', PA, assicurazioni).")

add_heading(doc, "7.1 Contesto di mercato", 2)
add_para(doc, "Molte aziende italiane strutturate (soprattutto settori regolamentati) hanno contratti enterprise Microsoft Azure, policy di data residency EU e processi di vendor security che impongono l'uso di cloud Microsoft. Queste aziende non possono adottare una SaaS \"off-the-shelf\" ospitata altrove.")

add_para(doc, "Al contempo, Anthropic Claude non e' nativamente disponibile su Azure (e' esposto via AWS Bedrock, Google Vertex e API diretta Anthropic). Le soluzioni tecniche sono due.")

build_table(
    doc,
    headers=["Path", "Architettura", "Vantaggi", "Svantaggi"],
    rows=[
        ["Path A - Ibrido", "App su Azure EU, Claude via AWS Bedrock EU Frankfurt", "Claude native quality, compliance EU", "Doppio cloud, latenza cross-region 40-80ms"],
        ["Path B - Azure-only", "App su Azure, AI via Azure OpenAI (GPT-5)", "Single-cloud, TCO ridotto", "Qualita' AI inferiore per copy marketing"],
    ],
    widths_cm=[3.0, 4.5, 4.0, 4.5],
)

add_heading(doc, "7.2 Pricing Azure Enterprise Edition", 2)
build_table(
    doc,
    headers=["Voce", "Importo", "Note"],
    rows=[
        ["Setup one-time", "€5.000", "Implementazione, data migration, training team"],
        ["Canone mensile base", "da €1.500", "Scala con seats e volume crediti"],
        ["Prezzo crediti AI", "€0,30/credito", "vs €0,03 standard (markup 10x per Azure/Bedrock)"],
        ["Seats dashboard aggiuntivi", "€25/seat/mese", "Oltre i 10 del piano base"],
        ["SLA 99,9% con penalty", "incluso", "Rimborso pro-rata in caso di downtime"],
        ["Support dedicato 24/7", "incluso", "Response SLA: 2h critical, 4h standard"],
        ["Security audit & reports", "incluso", "SOC2 + ISO 27001 annuale"],
    ],
    widths_cm=[5.5, 3.5, 7.0],
    highlight_rows={1: "F0FDFA", 4: "F0FDFA"},
)

add_heading(doc, "7.3 Analisi margini Azure Enterprise", 2)
add_para(doc, "Simulazione su tenant medio con canone €3.000/mese + 5.000 crediti AI consumati + 20 seats.")

build_table(
    doc,
    headers=["Voce", "Valore", "Note"],
    rows=[
        ["Revenue mensile", "€3.000", "Canone base €1.500 + crediti €1.500"],
        ["+ Revenue seats aggiuntivi", "€250", "10 seats extra @ €25"],
        ["Revenue totale", "€3.250", ""],
        ["Costo AI (5.000 crediti @ $0,034 su Bedrock)", "-€155", "Markup Bedrock +30% vs direct"],
        ["Costo infra Azure dedicata", "-€600", "AKS + Postgres Flexible + Azure Cache"],
        ["Costo CSM dedicato (quota tenant)", "-€400", "CSM condiviso su 10 tenant"],
        ["Costo compliance (audit, reports)", "-€150", "Quota annua spalmata"],
        ["Costo support dedicato (quota)", "-€250", ""],
        ["Costi totali", "-€1.555", ""],
        ["Margine lordo", "€1.695", "52,2%"],
        ["Margine annualizzato", "€20.340", ""],
    ],
    widths_cm=[7.0, 4.0, 5.0],
    highlight_rows={2: "F8FAFC", 8: "F8FAFC", 9: "F0FDFA", 10: "FFFBEB"},
)
add_callout(
    doc,
    "Il margine percentuale e' inferiore (52% vs 85%+ su Enterprise standard) ma la revenue assoluta "
    "e' 7-10x maggiore. LTV stimato per Azure Enterprise: €136.800 (36 mesi, retention 95%). "
    "Il CAC outbound (€5.000-15.000 via sales cycle 3-6 mesi) e' recuperato in 4-8 mesi."
)

add_heading(doc, "7.4 Setup team commerciale per Azure Enterprise", 2)
add_bullet(doc, "Lead qualification: aziende >50 dipendenti, settori regolamentati, contratto Azure attivo")
add_bullet(doc, "Sales cycle atteso: 3-6 mesi (technical review, compliance review, legal review, POC)")
add_bullet(doc, "Deal size medio: €40.000-80.000/anno (canone + consumo + setup)")
add_bullet(doc, "Materiale di supporto: case study, security whitepaper, SLA doc, DPA template")
add_bullet(doc, "Partner canale: system integrator Microsoft (Reply, Accenture, Engineering)")

doc.add_page_break()

# ===== 8. UNIT ECONOMICS & SCENARI =====
add_heading(doc, "8. Unit Economics & Scenari", 1)

add_heading(doc, "8.1 CAC target per canale", 2)
build_table(
    doc,
    headers=["Canale", "CAC atteso", "Target piano", "Payback"],
    rows=[
        ["Content marketing / SEO", "€50-150", "Starter + Free Trial", "1-3 mesi"],
        ["Paid ads (Google, Meta, LinkedIn)", "€200-400", "Professional", "2-4 mesi"],
        ["Outbound sales (cold email/LinkedIn)", "€800-1.500", "Enterprise", "3-5 mesi"],
        ["Partner channel (agenzie)", "€300-600", "Professional / Enterprise", "2-4 mesi"],
        ["Enterprise sales cycle Azure", "€5.000-15.000", "Azure Enterprise", "4-8 mesi"],
    ],
    widths_cm=[5.5, 3.0, 4.0, 3.5],
)

add_heading(doc, "8.2 LTV stimato", 2)
build_table(
    doc,
    headers=["Piano", "MRR medio", "Retention mesi", "Churn annuo", "LTV"],
    rows=[
        ["Starter", "€49", "12", "35%", "€353"],
        ["Professional", "€149", "18", "30%", "€1.880"],
        ["Enterprise", "€399", "24", "15%", "€8.140"],
        ["Azure Enterprise", "€3.800", "36", "5%", "€136.800"],
    ],
    widths_cm=[4.0, 3.0, 3.0, 3.0, 3.0],
    highlight_rows={3: "F0FDFA"},
)

add_heading(doc, "8.3 Break-even piattaforma", 2)
add_para(doc, "Costi fissi piattaforma stimati: €10.000/mese (include sviluppo prodotto, hosting, compliance, support, marketing operativo).")

build_table(
    doc,
    headers=["Scenario break-even", "Clienti richiesti", "ARR equivalente", "Fattibilita'"],
    rows=[
        ["Solo Starter", "230", "€135.000", "Molto difficile - bassa revenue/cliente"],
        ["Solo Professional", "73", "€130.000", "Fattibile in 6-9 mesi"],
        ["Solo Enterprise", "28", "€134.000", "Fattibile con outbound mirato"],
        ["Mix realistico Q1", "40 Pro + 8 Ent + 2 Azure", "€153.000", "Target piano marketing"],
    ],
    widths_cm=[5.0, 5.0, 3.5, 4.0],
    highlight_rows={3: "F0FDFA"},
)

add_heading(doc, "8.4 Scenari 12 mesi", 2)
build_table(
    doc,
    headers=["Scenario", "Starter", "Pro", "Ent", "Azure", "MRR mese 12", "ARR"],
    rows=[
        ["Conservativo", "50", "30", "5", "1", "€10.000", "€120.000"],
        ["Base", "100", "70", "15", "3", "€25.000", "€300.000"],
        ["Ambizioso", "200", "150", "40", "8", "€58.000", "€700.000"],
    ],
    widths_cm=[3.5, 2.2, 2.2, 2.2, 2.2, 3.0, 2.5],
    highlight_rows={1: "F0FDFA"},
)
add_para(doc, "Lo scenario \"Base\" e' quello usato per la pianificazione commerciale e di tesoreria del 2026. Richiede focus Pro in Q1-Q2 e spinta Enterprise in Q3-Q4.", 9, italic=True, color=GRAY_MUTED)

doc.add_page_break()

# ===== 9. GO-TO-MARKET =====
add_heading(doc, "9. Go-to-Market & Playbook Commerciale", 1)

add_heading(doc, "9.1 Segmenti target prioritari", 2)
build_table(
    doc,
    headers=["Segmento", "Piano tipico", "Use case", "Volume stimato IT"],
    rows=[
        ["Retail & E-commerce", "Professional", "Recupero carrello, upsell post-acquisto, nurturing", "~40.000 PMI"],
        ["Real estate", "Professional", "Follow-up lead, conferma appuntamenti", "~12.000 studi"],
        ["Healthcare privato", "Enterprise", "Conferma visite, follow-up prestazioni", "~8.000 cliniche"],
        ["Finanza / Assicurazioni", "Azure Enterprise", "Customer service, compliance notifications", "~3.500 operatori"],
        ["PA e Municipale", "Azure Enterprise", "Comunicazioni cittadini", "~7.900 enti"],
        ["Agenzie marketing", "Enterprise white-label", "Servizio ai clienti finali", "~3.000 agenzie"],
    ],
    widths_cm=[4.5, 3.5, 5.5, 3.0],
)

add_heading(doc, "9.2 Playbook commerciale", 2)
add_bullet(doc, "Touch 1: invio free trial via form sito (zero carta richiesta) - obiettivo 2.000 signup/trimestre")
add_bullet(doc, "Touch 2 (giorno 3): email follow-up con case study di settore - CTR target 25%")
add_bullet(doc, "Touch 3 (giorno 7): chiamata di onboarding 30min con sales - setup Twilio + primo template")
add_bullet(doc, "Touch 4 (giorno 12): email reminder scadenza trial - conversion target 20%")
add_bullet(doc, "Expansion: QBR (quarterly business review) per clienti >€300/mese - up-sell Enterprise / Azure")

add_heading(doc, "9.3 Marketing operativo", 2)
add_bullet(doc, "Content hub: tutorial settimanali su casi d'uso (AI campaigns, recupero carrello, reminder)")
add_bullet(doc, "SEO: parole chiave \"whatsapp business marketing\", \"ai campaigns italiano\", \"alternative mailchimp whatsapp\"")
add_bullet(doc, "Paid: LinkedIn Ads per Enterprise, Meta Ads per PMI, Google Search per intent query")
add_bullet(doc, "Partnership: accordo referral con 10 agenzie digitali, commissione 20% primi 12 mesi")
add_bullet(doc, "Eventi: Web Marketing Festival, SMAU, conferenze settore retail/real estate")

doc.add_page_break()

# ===== 10. ROADMAP COMMERCIALE =====
add_heading(doc, "10. Roadmap Commerciale 2026", 1)

build_table(
    doc,
    headers=["Periodo", "Obiettivi", "KPI target", "Owner"],
    rows=[
        ["Q1 (apr-giu)", "Lancio Pro + campagne paid IT", "30 Pro, 5 Ent, 5K signup trial", "Sales + Marketing"],
        ["Q2 (lug-set)", "Programma partner agenzie", "+20 Pro, +5 Ent, 10 agenzie attive", "Sales + Partnership"],
        ["Q3 (ott-dic)", "Lancio Azure Enterprise", "3 Azure tenant attivi", "Enterprise Sales"],
        ["Q4 (gen-mar 2027)", "Espansione Spagna/Germania", "First 10 clienti esteri", "International"],
    ],
    widths_cm=[3.0, 5.5, 4.5, 3.5],
)

add_heading(doc, "10.1 Budget marketing trimestre", 2)
build_table(
    doc,
    headers=["Voce", "Budget Q1", "Budget Q2", "Budget Q3", "Budget Q4"],
    rows=[
        ["Paid ads (Google + Meta + LinkedIn)", "€8.000", "€12.000", "€15.000", "€18.000"],
        ["Content / SEO (freelance + tools)", "€3.000", "€3.500", "€4.000", "€4.000"],
        ["Eventi + fiere", "€2.000", "€5.000", "€4.000", "€6.000"],
        ["Partnership / referral bonus", "€1.500", "€3.000", "€4.000", "€5.000"],
        ["Sales Enterprise outbound", "€2.000", "€4.000", "€8.000", "€10.000"],
        ["TOTALE", "€16.500", "€27.500", "€35.000", "€43.000"],
    ],
    widths_cm=[6.5, 2.5, 2.5, 2.5, 2.5],
    highlight_rows={5: "F0FDFA"},
)

doc.add_page_break()

# ===== 11. APPENDICE =====
add_heading(doc, "11. Appendice", 1)

add_heading(doc, "A. Glossario", 2)
add_bullet(doc, "BYOK - Bring Your Own Key. L'utente fornisce la sua API key Anthropic, paga Anthropic direttamente, Wamply non conteggia crediti.")
add_bullet(doc, "Credito AI - Unita' di misura dell'uso AI sulla piattaforma. 1 credito ~ $0,026 di costo Anthropic medio.")
add_bullet(doc, "Routing silente - Logica automatica che sceglie il modello Claude ottimale (Sonnet/Opus/Haiku) per ogni operazione. Invisibile all'utente.")
add_bullet(doc, "Top-up - Pacchetto di crediti one-shot acquistabile on demand, valido 12 mesi.")
add_bullet(doc, "CAC - Customer Acquisition Cost. Costo medio per acquisire un nuovo cliente pagante.")
add_bullet(doc, "LTV - Lifetime Value. Revenue totale attesa da un cliente nell'intera durata della relazione.")
add_bullet(doc, "ARR - Annual Recurring Revenue. Revenue ricorrente annualizzata dei clienti attivi.")
add_bullet(doc, "MRR - Monthly Recurring Revenue. Revenue ricorrente mensile.")

add_heading(doc, "B. Contatti & escalation", 2)
add_para(doc, "Per domande tecniche, contratti, proposte commerciali personalizzate e casi Azure Enterprise:")
add_para(doc, "Email commerciale: admin@agentengineering.it", bold=True)
add_para(doc, "Website: www.agentengineering.it")
add_para(doc, "Prodotto: www.wamply.com", italic=True)

add_heading(doc, "C. Note metodologiche", 2)
add_bullet(doc, "Tutti i dati economici sono in Euro salvo diversa indicazione.")
add_bullet(doc, "Cambio USD/EUR usato: 0,92 (aprile 2026).")
add_bullet(doc, "I costi Anthropic sono stimati sul pricing API standard aprile 2026 con mix d'uso tipico.")
add_bullet(doc, "I costi Twilio WhatsApp sono pass-through al cliente e non inclusi nei margini Wamply.")
add_bullet(doc, "Le percentuali di conversione e retention sono benchmark SaaS B2B, da validare nei primi 3-6 mesi.")
add_bullet(doc, "Il budget marketing e' soggetto a revisione trimestrale in base ai KPI di acquisizione.")

add_teal_divider(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run("Fine documento")
run.font.name = "Arial"
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = GRAY_MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Wamply  Business Case v1.0  21 aprile 2026")
run.font.name = "Arial"
run.font.size = Pt(9)
run.font.color.rgb = TEAL

# Save
out_path = OUT / "Wamply-Business-Case-v1.docx"
doc.save(str(out_path))
print(f"Saved: {out_path}")

import os
print(f"Size: {os.path.getsize(out_path) / 1024:.1f} KB")
